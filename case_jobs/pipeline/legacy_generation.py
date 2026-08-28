import asyncio
import json
import logging
import os
import re
import time
from dataclasses import dataclass

import markdown
from bs4 import BeautifulSoup
from bs4.element import NavigableString, Tag
from django.conf import settings
from docx import Document

from immigration_algoai_APIs.media_cleanup import clean_media_files_for_generation, delete_file_if_exists

from case_jobs.retrieval.visa_rag import get_or_build_corpus
from case_jobs.pipeline.context_enrichment import (
    ExhibitMarkdownAgent,
    sanitize_exhibits,
)


logger = logging.getLogger(__name__)

FENCED_BLOCK_PATTERN = re.compile(r"```(?:markdown)?\s*(.*?)```", re.DOTALL | re.IGNORECASE)


@dataclass(frozen=True)
class InlinePromptDocument:
    page_content: str
    metadata: dict


def build_submitted_parties_prompt(submitted_parties: dict) -> InlinePromptDocument:
    party_payload = {
        key: submitted_parties.get(key)
        for key in ("beneficiary", "petitioner", "preparer")
        if key in submitted_parties
    }
    payload = json.dumps(party_payload, sort_keys=True, separators=(",", ":"))
    content = "\n".join(
        (
            "# Submitted Parties (from request payload)",
            "Use supplied values as authoritative. Null values are unavailable and must not be invented.",
            payload,
        )
    ).strip()
    metadata = {
        "source_name": "request_payload",
        "source_category": "submitted_parties",
        "file_hash": "payload",
        "page_number": 0,
        "chunk_index": 0,
        "extraction_mode": "payload",
    }
    return InlinePromptDocument(page_content=content, metadata=metadata)


def build_exhibit_request_prompt(submitted_request: dict) -> InlinePromptDocument:
    beneficiary = submitted_request.get("beneficiary") or {}
    petitioner = submitted_request.get("petitioner") or {}
    compact_request = {
        "beneficiary_name": beneficiary.get("full_name"),
        "petitioner_name": petitioner.get("full_name"),
        "exhibits": sanitize_exhibits(submitted_request.get("exhibits", [])),
    }
    payload = json.dumps(compact_request, sort_keys=True, separators=(",", ":"))
    content = "\n".join(
        (
            "# Submitted Exhibit Request (from request payload)",
            "Use the visa-specific Exhibit List template and treat the compact payload below as authoritative.",
            "Preserve the exhibit array order, numbers, titles, item names, and item order exactly.",
            "Label each item as <exhibit number>.<one-based item position>.",
            "Include only supplied exhibits and items. Source URLs must not appear in the model input or final document; file contents are not required.",
            payload,
        )
    ).strip()
    return InlinePromptDocument(
        page_content=content,
        metadata={
            "source_name": "request_payload",
            "source_category": "submitted_request",
            "file_hash": "payload",
            "page_number": 0,
            "chunk_index": 0,
            "extraction_mode": "payload",
        },
    )


def build_inline_prompt(content: str, source_category: str) -> InlinePromptDocument:
    return InlinePromptDocument(
        page_content=content.strip(),
        metadata={
            "source_name": source_category,
            "source_category": source_category,
            "file_hash": source_category,
            "page_number": 0,
            "chunk_index": 0,
            "extraction_mode": "generated_context",
        },
    )


def _preview_text(value: str, limit: int) -> str:
    if not value:
        return ""
    if limit <= 0:
        return ""
    value = value.strip()
    if len(value) <= limit:
        return value
    return f"{value[:limit].rstrip()}…"


def extract_markdown_output(response_text: str) -> str:
    matches = FENCED_BLOCK_PATTERN.findall((response_text or "").strip())
    if matches:
        return "\n\n".join(match.strip() for match in matches if match.strip()).strip()
    return (response_text or "").strip()

def _paragraph_ends_with_whitespace(paragraph) -> bool:
    runs = getattr(paragraph, "runs", None) or ()
    if not runs:
        return True
    text = runs[-1].text or ""
    return not text or text[-1].isspace()


def _append_text_with_breaks(paragraph, text: str, *, bold: bool = False) -> None:
    if text is None:
        return
    normalized = str(text).replace("\r\n", "\n").replace("\r", "\n")
    parts = normalized.split("\n")
    for index, part in enumerate(parts):
        if part:
            run = paragraph.add_run(part)
            if bold:
                run.bold = True
        if index < len(parts) - 1:
            paragraph.add_run().add_break()


def _append_inline_node(paragraph, node) -> None:
    if node is None:
        return
    if isinstance(node, Tag):
        if node.name == "br":
            paragraph.add_run().add_break()
            return
        if node.name == "strong":
            _append_text_with_breaks(paragraph, node.get_text(), bold=True)
            return
        _append_text_with_breaks(paragraph, node.get_text())
        return
    if isinstance(node, NavigableString):
        raw = str(node)
        if not raw:
            return
        if raw.isspace():
            if not _paragraph_ends_with_whitespace(paragraph):
                paragraph.add_run(" ")
            return
        _append_text_with_breaks(paragraph, raw)
        return


async def convert_markdown_to_docx(md_text: str, output_path: str) -> str | None:
    """
    Convert Markdown text to a .docx file asynchronously and save it in MEDIA_ROOT/generated.
    """
    try:
        generated_dir = os.path.join(settings.MEDIA_ROOT, "generated")
        os.makedirs(generated_dir, exist_ok=True)
        full_output_path = os.path.join(generated_dir, output_path)
        delete_file_if_exists(full_output_path)

        html = markdown.markdown(md_text, extensions=["extra", "nl2br"])
        soup = BeautifulSoup(html, "html.parser")

        document = Document()
        for element in soup.children:
            if not isinstance(element, Tag):
                if isinstance(element, str) and element.strip():
                    document.add_paragraph(element.strip())
                continue
            if element.name == "h1":
                document.add_heading(element.get_text(), level=1)
            elif element.name == "h2":
                document.add_heading(element.get_text(), level=2)
            elif element.name == "h3":
                document.add_heading(element.get_text(), level=3)
            elif element.name == "p":
                paragraph = document.add_paragraph()
                for node in element.children:
                    _append_inline_node(paragraph, node)
            elif element.name == "hr":
                document.add_page_break()
            elif element.name == "table":
                rows = element.find_all("tr")
                column_count = max(
                    (len(row.find_all(["th", "td"], recursive=False)) for row in rows),
                    default=0,
                )
                if column_count:
                    table = document.add_table(rows=len(rows), cols=column_count)
                    table.style = "Table Grid"
                    for row_index, source_row in enumerate(rows):
                        source_cells = source_row.find_all(["th", "td"], recursive=False)
                        for column_index, source_cell in enumerate(source_cells):
                            paragraph = table.rows[row_index].cells[column_index].paragraphs[0]
                            if source_cell.name == "th":
                                paragraph.add_run(source_cell.get_text(" ", strip=True)).bold = True
                            else:
                                for node in source_cell.children:
                                    _append_inline_node(paragraph, node)
            elif element.name == "ul":
                for list_item in element.find_all("li", recursive=False):
                    document.add_paragraph(list_item.get_text(), style="ListBullet")
            elif element.name == "ol":
                for list_item in element.find_all("li", recursive=False):
                    document.add_paragraph(list_item.get_text(), style="ListNumber")

        document.save(full_output_path)
        logger.info("Saved L1a DOCX output to %s", full_output_path)
        return full_output_path
    except Exception:
        logger.exception("Error converting Markdown to .docx for %s", output_path)
        return None


def summarise_retrieved_chunks(retrieved_documents) -> list[str]:
    source_names: list[str] = []
    for document in retrieved_documents:
        metadata = getattr(document, "metadata", {}) or {}
        source_names.append(str(metadata.get("source_name", "unknown")))
    return source_names


async def async_handle_doc_generation(
    file_manifests: list[dict],
    selected_options: list[str],
    adapter,
    *,
    tenant_id: str = "legacy",
    case_id: str = "legacy",
    submitted_parties: dict | None = None,
    corrections: list[str] | None = None,
) -> list[str]:
    """
    Generate documents based on downloaded source files and selected output options.
    """
    total_start = time.perf_counter()

    try:
        clean_media_files_for_generation()

        submitted_exhibits = (
            submitted_parties.get("exhibits", [])
            if isinstance(submitted_parties, dict)
            else []
        )
        payload_only_exhibit_list = bool(submitted_exhibits) and bool(selected_options) and all(
            option == "Exhibit List" for option in selected_options
        )
        if payload_only_exhibit_list:
            corpus_bundle = None
            processed_sources: list[dict] = []
            corpus_hash = "payload-only"
            cache_hit = False
            chunk_count = 0
            logger.info("Skipping corpus ingestion for structured Exhibit List")
        else:
            ingest_start = time.perf_counter()
            corpus_bundle = await get_or_build_corpus(
                file_manifests,
                config=adapter.rag_config(tenant_id=tenant_id, case_id=case_id),
            )
            processed_sources = corpus_bundle.processed_sources
            corpus_hash = corpus_bundle.corpus_hash
            cache_hit = corpus_bundle.cache_hit
            chunk_count = corpus_bundle.chunk_count
            logger.info(
                "Corpus %s ready in %.2fs | cache_hit=%s | chunks=%s",
                corpus_hash,
                time.perf_counter() - ingest_start,
                cache_hit,
                chunk_count,
            )
            for source in processed_sources:
                logger.info(
                    "Processed source %s | mode=%s | pages=%s | latex=%s",
                    source.get("original_filename"),
                    source.get("extraction_mode"),
                    source.get("page_count"),
                    source.get("latex_path"),
                )

        generated_files: list[str] = []
        submitted_prompt = (
            build_submitted_parties_prompt(submitted_parties)
            if isinstance(submitted_parties, dict) and submitted_parties
            else None
        )
        exhibit_markdown = ""
        if submitted_exhibits and any(
            option != "Exhibit List" for option in selected_options
        ):
            exhibit_markdown = await ExhibitMarkdownAgent().generate(submitted_exhibits)

        for file_type in selected_options:
            retrieval_start = time.perf_counter()
            retrieved_documents = (
                []
                if payload_only_exhibit_list or corpus_bundle is None
                else [
                    InlinePromptDocument(
                        page_content=chunk.text,
                        metadata={
                            "source_name": chunk.metadata.get("source_name", "unknown")
                        },
                    )
                    for chunk in corpus_bundle.chunks
                ]
            )
            context_documents: list = []
            if corrections:
                correction_text = "\n".join(
                    ["# Verification Corrections", "Apply every correction below:"]
                    + [f"- {correction}" for correction in corrections]
                )
                context_documents.append(
                    build_inline_prompt(correction_text, "verification_corrections")
                )

            if file_type == "Exhibit List":
                if submitted_exhibits and isinstance(submitted_parties, dict):
                    context_documents.append(
                        build_exhibit_request_prompt(submitted_parties)
                    )
                elif submitted_prompt is not None:
                    context_documents.append(submitted_prompt)
            else:
                if submitted_prompt is not None:
                    context_documents.append(submitted_prompt)
                if exhibit_markdown:
                    context_documents.append(
                        build_inline_prompt(exhibit_markdown, "submitted_exhibits_markdown")
                    )

            retrieved_documents = [*context_documents, *retrieved_documents]
            logger.info(
                "Added all %s corpus chunks for %s in %.2fs | source_names=%s",
                len(retrieved_documents),
                file_type,
                time.perf_counter() - retrieval_start,
                summarise_retrieved_chunks(retrieved_documents),
            )

            generation_start = time.perf_counter()
            agent_response = await adapter.generate_document(
                file_type=file_type,
                retrieved_context=retrieved_documents,
                source_manifest=[],
            )
            logger.info(
                "Generated %s response in %.2fs",
                file_type,
                time.perf_counter() - generation_start,
            )

            if not agent_response:
                logger.warning("No response returned for %s", file_type)
                continue

            markdown_text = extract_markdown_output(agent_response)
            should_print_markdown = bool(getattr(settings, "PRINT_GENERATION_MARKDOWN", False))
            if markdown_text and should_print_markdown:
                print(
                    "\n".join(
                        (
                            f"=== generated_markdown visa_type={getattr(adapter, 'visa_type', 'unknown')} file_type={file_type} ===",
                            markdown_text,
                            "=== end_generated_markdown ===",
                        )
                    ),
                    flush=True,
                )
            preview_limit = int(getattr(settings, "GENERATION_LOG_PREVIEW_CHARS", 2000) or 2000)
            if markdown_text:
                logger.info(
                    "generation markdown preview visa_type=%s file_type=%s corpus_hash=%s cache_hit=%s chunk_count=%s markdown_len=%s preview=%r",
                    getattr(adapter, "visa_type", "unknown"),
                    file_type,
                    corpus_hash,
                    cache_hit,
                    chunk_count,
                    len(markdown_text),
                    _preview_text(markdown_text, preview_limit),
                )
            else:
                logger.warning(
                    "generation produced no markdown visa_type=%s file_type=%s corpus_hash=%s cache_hit=%s chunk_count=%s response_len=%s response_preview=%r",
                    getattr(adapter, "visa_type", "unknown"),
                    file_type,
                    corpus_hash,
                    cache_hit,
                    chunk_count,
                    len(agent_response),
                    _preview_text(agent_response, preview_limit),
                )
            if not markdown_text:
                logger.warning("No Markdown content found for %s", file_type)
                continue

            output_filename = f"{file_type.replace(' ', '_')}.docx"
            generated_file_path = await convert_markdown_to_docx(markdown_text, output_filename)
            if generated_file_path:
                generated_files.append(generated_file_path)

        logger.info(
            "%s generation completed in %.2fs",
            adapter.display_name,
            time.perf_counter() - total_start,
        )
        return generated_files
    except Exception:
        logger.exception("Error in %s document generation pipeline.", adapter.display_name)
        return []


def handle_doc_generation(
    file_manifests: list[dict],
    selected_options: list[str],
    adapter,
    *,
    tenant_id: str = "legacy",
    case_id: str = "legacy",
    submitted_parties: dict | None = None,
    corrections: list[str] | None = None,
) -> list[str]:
    """
    Synchronous wrapper that runs the async document generation pipeline.
    """
    return asyncio.run(
        async_handle_doc_generation(
            file_manifests,
            selected_options,
            adapter,
            tenant_id=tenant_id,
            case_id=case_id,
            submitted_parties=submitted_parties,
            corrections=corrections,
        )
    )
