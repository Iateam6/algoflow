import os
from PIL import Image, ImageDraw
from PyPDF2 import PdfMerger
from docx import Document
from copy import deepcopy
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Pt
from docx.enum.table import WD_ALIGN_VERTICAL, WD_ROW_HEIGHT_RULE
from docx.enum.text import WD_BREAK, WD_PARAGRAPH_ALIGNMENT
from docxcompose.composer import Composer
from pdfminer.pdfparser import PDFParser
from pdfminer.pdfdocument import PDFDocument
from pdfminer.pdfparser import PDFSyntaxError
from reportlab.pdfgen import canvas
from reportlab.lib.pagesizes import letter

import os
from pathlib import Path
from copy import deepcopy
from typing import List, Tuple, Union, Optional, Any, cast
import shutil
import subprocess

from io import BytesIO
from reportlab.pdfbase import pdfmetrics
from pypdf import PdfReader, PdfWriter

async def convert_to_pdf(file_path):
    """Convert image to PDF."""
    try:
        pdf_path = os.path.splitext(file_path)[0] + ".pdf"
        with Image.open(file_path) as img:
            img = img.convert("RGB")
            img.save(pdf_path, "PDF", resolution=100.0)
        print(f"[OK] Converted to PDF: {pdf_path}")
        return pdf_path
    except Exception as e:
        print(f"[ERROR] Image conversion failed for {file_path}: {e}")
        return file_path


async def merge_pdfs(pdf_paths, output_path):
    """
    Merge PDFs safely across Windows & Linux.
    - Validates PDFs before merging
    - Automatically skips or replaces corrupted files
    """
    merger = PdfMerger()
    valid_pdfs = []

    for path in pdf_paths:
        if not path or not os.path.exists(path):
            print(f"[WARN] Skipping missing file: {path}")
            continue

        # Validate PDF with pdfminer.six
        try:
            with open(path, "rb") as f:
                parser = PDFParser(f)
                PDFDocument(parser)  # attempt to parse
            valid_pdfs.append(path)
            print(f"[OK] Valid PDF: {path}")

        except (PDFSyntaxError, Exception) as e:
            print(f"[WARN] Invalid PDF detected: {path} ({e})")
            placeholder = output_path + f"_{os.path.basename(path)}_error.pdf"
            c = canvas.Canvas(placeholder, pagesize=letter)
            c.setFont("Helvetica", 12)
            c.drawCentredString(300, 500, f"⚠️ Skipped file: {os.path.basename(path)}")
            c.drawCentredString(300, 470, f"Error: {str(e)[:100]}")
            c.save()
            valid_pdfs.append(placeholder)
            print(f"[INFO] Added placeholder for bad file: {placeholder}")

    if not valid_pdfs:
        raise Exception("No valid PDF files to merge.")

    try:
        for pdf in valid_pdfs:
            merger.append(pdf)
        merger.write(output_path)
        merger.close()
        print(f"[OK] Successfully merged {len(valid_pdfs)} PDFs → {output_path}")
    except Exception as e:
        print(f"[ERROR] PDF merge failed: {e}")
        raise


async def create_blank_page_pdf(output_path, text=""):
    """Create a blank separator page with centered text."""
    c = canvas.Canvas(output_path, pagesize=letter)
    width, height = letter

    if text:
        lines = [ln.strip() for ln in str(text).splitlines() if ln.strip()]
        if not lines:
            lines = [str(text)]

        font_name = "Helvetica-Bold"
        font_size = 16
        leading = int(font_size * 1.6)
        c.setFont(font_name, font_size)

        block_height = leading * (len(lines) - 1)
        start_y = (height / 2) + (block_height / 2)
        for idx, line in enumerate(lines):
            c.drawCentredString(width / 2, start_y - (idx * leading), line)

    c.showPage()
    c.save()
    print(f"[OK] Created separator PDF: {output_path}")


def stamp_pdf_with_firm_name(input_pdf_path: str, firm_name: str, output_pdf_path: str | None = None) -> str:
    """
    Stamp `firm_name` onto every page of `input_pdf_path`, returning the stamped PDF path.

    Fail-open: if stamping fails for any reason, returns `input_pdf_path`.
    """
    try:
        if not input_pdf_path or not os.path.exists(input_pdf_path):
            return input_pdf_path

        firm_name = (firm_name or "").strip()
        if not firm_name:
            return input_pdf_path

        base, ext = os.path.splitext(input_pdf_path)
        if ext.lower() != ".pdf":
            return input_pdf_path

        out_path = output_pdf_path or f"{base}_stamped.pdf"
        if out_path == input_pdf_path:
            out_path = f"{base}_stamped.pdf"

        # If already created in this temp dir run, reuse it.
        try:
            if os.path.exists(out_path) and os.path.getsize(out_path) > 0:
                return out_path
        except Exception:
            pass

        reader = PdfReader(input_pdf_path)
        if getattr(reader, "is_encrypted", False):
            try:
                reader.decrypt("")  # best-effort
            except Exception:
                print(f"[WARN] PDF is encrypted; skipping firm-name stamp: {os.path.basename(input_pdf_path)}")
                return input_pdf_path

        writer = PdfWriter()

        max_font_size = 9
        min_font_size = 6
        top_padding = 24  # points from top edge of visible box
        side_margin = 36  # points

        def _fit_text(text: str, visible_width: float) -> tuple[str, int]:
            max_width = max(0.0, float(visible_width) - (2 * side_margin))
            font_name = "Helvetica-Bold"

            font_size = max_font_size
            while font_size >= min_font_size:
                if pdfmetrics.stringWidth(text, font_name, font_size) <= max_width:
                    return text, font_size
                font_size -= 1

            # Still too long at min font size; truncate with ellipsis.
            font_size = min_font_size
            ellipsis = "..."
            if pdfmetrics.stringWidth(text, font_name, font_size) <= max_width:
                return text, font_size
            if pdfmetrics.stringWidth(ellipsis, font_name, font_size) > max_width:
                return "", font_size

            trimmed = text
            while trimmed and pdfmetrics.stringWidth(trimmed + ellipsis, font_name, font_size) > max_width:
                trimmed = trimmed[:-1]
            return (trimmed + ellipsis) if trimmed else "", font_size

        for page in reader.pages:
            try:
                # Use the visible box (CropBox) when present so the stamp lands in the viewer-visible area.
                media = page.mediabox
                crop = getattr(page, "cropbox", None) or media

                media_llx = float(media.lower_left[0])
                media_lly = float(media.lower_left[1])
                media_urx = float(media.upper_right[0])
                media_ury = float(media.upper_right[1])
                media_w = float(media_urx - media_llx)
                media_h = float(media_ury - media_lly)

                crop_llx = float(crop.lower_left[0])
                crop_lly = float(crop.lower_left[1])
                crop_urx = float(crop.upper_right[0])
                crop_ury = float(crop.upper_right[1])
                visible_w = float(crop_urx - crop_llx)

                # Desired stamp point in PDF user-space coordinates (relative to page boxes).
                stamp_x = crop_llx + (visible_w / 2.0)
                stamp_y = crop_ury - top_padding

                # ReportLab uses a (0,0) origin; shift PDF coords into overlay coords
                # where (0,0) corresponds to the MediaBox lower-left.
                overlay_x = stamp_x - media_llx
                overlay_y = stamp_y - media_lly

                draw_text, font_size = _fit_text(firm_name, visible_w)
                overlay_buf = BytesIO()
                c = canvas.Canvas(overlay_buf, pagesize=(media_w, media_h))
                if draw_text:
                    c.setFillColorRGB(0, 0, 0)
                    c.setFont("Helvetica-Bold", font_size)
                    c.drawCentredString(overlay_x, overlay_y, draw_text)
                c.showPage()
                c.save()
                overlay_buf.seek(0)

                overlay_reader = PdfReader(overlay_buf)
                overlay_page = overlay_reader.pages[0]
                page.merge_page(overlay_page)
                writer.add_page(page)
            except Exception as e:
                print(f"[WARN] Failed to stamp page; leaving as-is: {e}")
                writer.add_page(page)

        with open(out_path, "wb") as f:
            writer.write(f)

        return out_path
    except Exception as e:
        print(f"[WARN] Failed to stamp PDF with firm name: {e}")
        return input_pdf_path


def convert_docx_to_pdf(docx_path: str, output_dir: str | None = None) -> str | None:
    """
    Convert a local Word file (.docx or .doc) to .pdf and return the resulting PDF path.

    Implementation prefers LibreOffice (`soffice`) headless conversion.
    Returns None when conversion is unavailable or fails.
    """
    if not docx_path or not os.path.exists(docx_path):
        return None
    if os.path.splitext(docx_path)[1].lower() not in (".docx", ".doc"):
        return None

    outdir = output_dir or os.path.dirname(docx_path)
    os.makedirs(outdir, exist_ok=True)

    soffice = shutil.which("soffice") or shutil.which("libreoffice")
    if not soffice:
        print("[WARN] Word->PDF conversion unavailable (LibreOffice not found).")
        return None

    try:
        subprocess.run(
            [soffice, "--headless", "--nologo", "--nolockcheck", "--convert-to", "pdf", "--outdir", outdir, docx_path],
            check=True,
            stdout=subprocess.PIPE,
            stderr=subprocess.PIPE,
        )
    except Exception as e:
        print(f"[WARN] Word->PDF conversion failed for {docx_path}: {e}")
        return None

    expected_pdf = os.path.join(outdir, f"{Path(docx_path).stem}.pdf")
    if os.path.exists(expected_pdf) and os.path.getsize(expected_pdf) > 0:
        print(f"[OK] Converted Word to PDF: {expected_pdf}")
        return expected_pdf

    candidates = list(Path(outdir).glob(f"{Path(docx_path).stem}*.pdf"))
    for c in candidates:
        try:
            if c.exists() and c.stat().st_size > 0:
                print(f"[OK] Converted Word to PDF: {c}")
                return str(c)
        except Exception:
            continue

    print(f"[WARN] Word->PDF conversion produced no output for {docx_path}")
    return None


# DOCX utilities

def _is_sectPr(element) -> bool:
    return element.tag.endswith('}sectPr') or element.tag.endswith('sectPr')


def _normalize(doc_entries):
    normalized = []
    for e in doc_entries:
        if isinstance(e, (list, tuple)) and len(e) >= 2:
            normalized.append((str(e[0]), str(e[1])))
        else:
            p = str(e)
            normalized.append((os.path.basename(p), p))
    return normalized


def prepend_cover_and_merge(
   doc_entries: List[Union[str, Tuple[str, str]]],
    merged_output_path: str,
    name_font_size: int = 16,
) -> str:
    """
    Each file -> make a copy with a true full-page cover (its own filename centered both
    vertically & horizontally).  Those new files are saved in 'with_covers' subfolder.
    Then merge all new files (once each, in order) into merged_output_path.
    """
    normalized = _normalize(doc_entries)
    merged_output = Path(merged_output_path)
    output_dir = merged_output.parent / "with_covers"
    output_dir.mkdir(parents=True, exist_ok=True)

    new_files = []

    # Step 1: create per-file cover + content
    for display_name, path in normalized:
        p = Path(path)
        if not p.exists() or not p.suffix.lower().endswith(".docx"):
            print(f"[SKIP] {path}")
            continue

        try:
            original = Document(str(p))
            new_doc = Document()

            # Copy section layout
            try:
                src = original.sections[0]
                dst = new_doc.sections[0]
                dst.page_width, dst.page_height = src.page_width, src.page_height
                dst.top_margin, dst.bottom_margin = src.top_margin, src.bottom_margin
                dst.left_margin, dst.right_margin = src.left_margin, src.right_margin
            except Exception:
                pass

            # Remove default empty paragraph
            if new_doc.paragraphs:
                p0 = new_doc.paragraphs[0]
                p0._element.getparent().remove(p0._element)

            # Printable area
            sec = new_doc.sections[0]
            printable_w = None
            printable_h = None

            page_width = sec.page_width
            page_height = sec.page_height
            top_margin = sec.top_margin
            bottom_margin = sec.bottom_margin
            left_margin = sec.left_margin
            right_margin = sec.right_margin

            if page_width is not None and left_margin is not None and right_margin is not None:
                printable_w = page_width - left_margin - right_margin
            if page_height is not None and top_margin is not None and bottom_margin is not None:
                printable_h = page_height - top_margin - bottom_margin

            # 1x1 table for true centered text
            tbl = new_doc.add_table(rows=1, cols=1)
            tbl.autofit = False
            cell = tbl.cell(0, 0)
            cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
            para = cell.paragraphs[0]
            para.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
            run = para.add_run(display_name)
            run.bold = True
            run.font.size = Pt(name_font_size)

            # borderless
            tbl_pr = tbl._tblPr
            borders = OxmlElement("w:tblBorders")
            for side in ("top", "left", "bottom", "right", "insideH", "insideV"):
                b = OxmlElement(f"w:{side}")
                b.set(qn("w:val"), "nil")
                borders.append(b)
            tbl_pr.append(borders)

            # size table
            try:
                if printable_w is not None:
                    cast(Any, tbl.columns[0]).width = printable_w
                row = tbl.rows[0]
                if printable_h is not None:
                    cast(Any, row).height = printable_h
                cast(Any, row).height_rule = WD_ROW_HEIGHT_RULE.EXACTLY
            except Exception:
                pass

            para.add_run().add_break(WD_BREAK.PAGE)

            # append original body
            for el in original.element.body:
                new_doc.element.body.append(deepcopy(el))

            # save in new folder
            new_path = output_dir / p.name
            new_doc.save(str(new_path))
            new_files.append(new_path)
            print(f"[OK] Created cover for {p.name}")

        except Exception as e:
            print(f"[ERROR] {path}: {e}")

    # Step 2: merge only those new files (once each)
    merged = Document()
    if merged.paragraphs:
        p0 = merged.paragraphs[0]
        p0._element.getparent().remove(p0._element)

    for nf in new_files:
        try:
            doc = Document(str(nf))
            for el in doc.element.body:
                if not _is_sectPr(el):
                    merged.element.body.append(deepcopy(el))
            print(f"[MERGED] {nf.name}")
        except Exception as e:
            print(f"[ERROR MERGING] {nf}: {e}")

    merged.save(str(merged_output))
    print(f"[DONE] Merged DOCX saved: {merged_output}")
    return str(merged_output)
