from __future__ import annotations

import json
import re
from dataclasses import dataclass

from django.conf import settings
from docx import Document
from openai import OpenAI

from case_jobs.api.schemas import GenerationRequest
from case_jobs.pipeline.identity import IdentityRecord


BRACKETED_FIELD_PATTERN = re.compile(r"\[([^\[\]\r\n]{2,120})\]")
ALLOWED_REVIEW_MARKER_PREFIXES = ("missing:", "review:", "stop:")


def contains_unresolved_template_field(text: str) -> bool:
    for match in BRACKETED_FIELD_PATTERN.finditer(text or ""):
        field = match.group(1).strip()
        normalized = field.casefold()
        if normalized.startswith(ALLOWED_REVIEW_MARKER_PREFIXES):
            continue
        if normalized == "exhibit — not provided":
            continue
        if field.isdigit():
            continue
        return True
    return False


@dataclass(frozen=True)
class VerificationResult:
    passed: bool
    corrections: tuple[str, ...] = ()


def extract_docx_text(path: str) -> str:
    document = Document(path)
    blocks = [paragraph.text for paragraph in document.paragraphs]
    for table in document.tables:
        for row in table.rows:
            blocks.append(" | ".join(cell.text for cell in row.cells))
    return "\n".join(blocks).strip()


def deterministic_verify(text: str, request: GenerationRequest) -> VerificationResult:
    corrections: list[str] = []
    lowered = text.lower()
    is_exhibit_list = bool(request.exhibits) and (
        request.document_type == "Exhibit List"
        or request.document_slug == "exhibit-list"
    )
    required_parties = [
        ("beneficiary", request.beneficiary.full_name),
        ("petitioner", request.petitioner.full_name),
    ]
    if not is_exhibit_list:
        required_parties.append(("preparer", request.preparer.full_name))
    for role, name in required_parties:
        if name and name.lower() not in lowered:
            corrections.append(f"Include the verified {role} name: {name}")
    if is_exhibit_list and request.exhibits:
        cursor = 0
        for exhibit in request.exhibits:
            if exhibit.title and exhibit.title.lower() not in lowered:
                corrections.append(
                    f"Include the submitted Exhibit {exhibit.number} title: {exhibit.title}"
                )
            for item_index, item in enumerate(exhibit.items, start=1):
                label = f"{exhibit.number}.{item_index}"
                label_position = lowered.find(label, cursor)
                if label_position < 0:
                    corrections.append(
                        f"Include exhibit item {label} in the submitted order: {item.name}"
                    )
                    continue
                name_position = lowered.find(item.name.lower(), label_position)
                if name_position < 0:
                    corrections.append(f"Include the submitted item name for {label}: {item.name}")
                    continue
                cursor = name_position + len(item.name)
    for source in request.files:
        if source.url.lower() in lowered:
            corrections.append(
                "Remove source URLs from the exhibit list"
                if is_exhibit_list
                else "Remove source URLs from the generated document"
            )
            break
    if contains_unresolved_template_field(text):
        corrections.append("Remove unresolved template placeholders")
    if len(text.split()) < 50:
        corrections.append("The generated document is unexpectedly short")
    return VerificationResult(not corrections, tuple(corrections))


class VerificationAgent:
    def __init__(self, client=None):
        self.client = client

    def verify(
        self,
        docx_path: str,
        request: GenerationRequest,
        identity: IdentityRecord | None,
        evidence_text: str,
    ) -> VerificationResult:
        draft = extract_docx_text(docx_path)
        deterministic = deterministic_verify(draft, request)
        is_structured_exhibit_list = bool(request.exhibits) and (
            request.document_type == "Exhibit List"
            or request.document_slug == "exhibit-list"
        )
        if is_structured_exhibit_list:
            return deterministic
        if not deterministic.passed or not settings.ENABLE_MODEL_VERIFICATION:
            return deterministic
        if identity is None:
            return VerificationResult(False, ("Verified identity is unavailable",))

        client = self.client or OpenAI(api_key=settings.OPENAI_API_KEY)
        prompt = {
            "instruction": (
                "Verify the generated immigration document against the submitted "
                "identity and evidence. Return only JSON with passed (boolean) and "
                "corrections (array of concise strings). Reject unsupported facts, "
                "wrong people, wrong roles, addresses, raw template placeholders, or "
                "missing sections. Explicit [MISSING: ...], [REVIEW: ...], and [STOP: ...] "
                "markers are permitted review indicators and are not raw placeholders."
            ),
            "submitted": request.to_dict(),
            "verified_identity": identity.to_dict(),
            "evidence_excerpt": evidence_text[:30000],
            "draft": draft[:50000],
        }
        try:
            response = client.responses.create(
                model=settings.VERIFICATION_MODEL,
                input=json.dumps(prompt, separators=(",", ":")),
            )
            raw = (response.output_text or "").strip()
            if raw.startswith("```"):
                raw = re.sub(r"^```(?:json)?\s*|\s*```$", "", raw)
            result = json.loads(raw)
            corrections = tuple(str(item) for item in result.get("corrections", []))
            return VerificationResult(bool(result.get("passed")) and not corrections, corrections)
        except Exception:
            return VerificationResult(False, ("Verification agent did not return a valid result",))
