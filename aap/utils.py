import asyncio
import logging
import os
import re
import time

import markdown
from bs4 import BeautifulSoup
from django.conf import settings
from docx import Document

from .agent import build_retrieval_query, generate_document
from .rag import get_or_build_corpus


logger = logging.getLogger(__name__)

FENCED_BLOCK_PATTERN = re.compile(r"```(?:markdown)?\s*(.*?)```", re.DOTALL | re.IGNORECASE)


def extract_markdown_output(response_text: str) -> str:
    matches = FENCED_BLOCK_PATTERN.findall((response_text or "").strip())
    if matches:
        return "\n\n".join(match.strip() for match in matches if match.strip()).strip()
    return (response_text or "").strip()


async def convert_markdown_to_docx(md_text: str, output_path: str) -> str | None:
    """
    Convert Markdown text to a .docx file asynchronously and save it in MEDIA_ROOT/generated.
    """
    try:
        generated_dir = os.path.join(settings.MEDIA_ROOT, "generated")
        os.makedirs(generated_dir, exist_ok=True)
        full_output_path = os.path.join(generated_dir, output_path)

        html = markdown.markdown(md_text, extensions=["extra"])
        soup = BeautifulSoup(html, "html.parser")

        document = Document()
        for element in soup.children:
            if element.name == "h1":
                document.add_heading(element.get_text(), level=1)
            elif element.name == "h2":
                document.add_heading(element.get_text(), level=2)
            elif element.name == "h3":
                document.add_heading(element.get_text(), level=3)
            elif element.name == "p":
                paragraph = document.add_paragraph()
                for node in element.children:
                    text = node.get_text() if hasattr(node, "get_text") else str(node)
                    if node.name == "strong":
                        run = paragraph.add_run(text)
                        run.bold = True
                    else:
                        paragraph.add_run(text)
            elif element.name == "hr":
                document.add_page_break()
            elif element.name == "ul":
                for list_item in element.find_all("li", recursive=False):
                    document.add_paragraph(list_item.get_text(), style="ListBullet")
            elif element.name == "ol":
                for list_item in element.find_all("li", recursive=False):
                    document.add_paragraph(list_item.get_text(), style="ListNumber")
            elif isinstance(element, str) and element.strip():
                document.add_paragraph(element.strip())

        document.save(full_output_path)
        logger.info("Saved L1a DOCX output to %s", full_output_path)
        return full_output_path
    except Exception:
        logger.exception("Error converting Markdown to .docx for %s", output_path)
        return None


def summarise_retrieved_chunks(retrieved_documents) -> list[str]:
    chunk_ids: list[str] = []
    for document in retrieved_documents:
        metadata = getattr(document, "metadata", {}) or {}
        chunk_ids.append(
            ":".join(
                [
                    str(metadata.get("file_hash", "unknown")),
                    str(metadata.get("page_number", "unknown")),
                    str(metadata.get("chunk_index", "unknown")),
                ]
            )
        )
    return chunk_ids


async def async_handle_doc_generation(
    file_manifests: list[dict],
    selected_options: list[str],
) -> list[str]:
    """
    Generate documents based on downloaded source files and selected output options.
    """
    total_start = time.perf_counter()

    try:
        ingest_start = time.perf_counter()
        corpus_bundle = await get_or_build_corpus(file_manifests)
        logger.info(
            "Corpus %s ready in %.2fs | cache_hit=%s | chunks=%s",
            corpus_bundle.corpus_hash,
            time.perf_counter() - ingest_start,
            corpus_bundle.cache_hit,
            corpus_bundle.chunk_count,
        )

        for source in corpus_bundle.processed_sources:
            logger.info(
                "Processed source %s | mode=%s | pages=%s | latex=%s",
                source.get("original_filename"),
                source.get("extraction_mode"),
                source.get("page_count"),
                source.get("latex_path"),
            )

        generated_files: list[str] = []
        for file_type in selected_options:
            retrieval_query = build_retrieval_query(file_type)
            retrieval_start = time.perf_counter()
            retrieved_documents = corpus_bundle.retrieve(retrieval_query)
            logger.info(
                "Retrieved %s chunks for %s in %.2fs | chunk_ids=%s",
                len(retrieved_documents),
                file_type,
                time.perf_counter() - retrieval_start,
                summarise_retrieved_chunks(retrieved_documents),
            )

            generation_start = time.perf_counter()
            agent_response = await generate_document(
                file_type=file_type,
                retrieved_context=retrieved_documents,
                source_manifest=corpus_bundle.processed_sources,
            )
            logger.info(
                "Generated %s response in %.2fs",
                file_type,
                time.perf_counter() - generation_start,
            )

            if not agent_response:
                logger.warning("No response returned for %s", file_type)
                continue

            markdown_text = extract_markdown_output(agent_response)
            if not markdown_text:
                logger.warning("No Markdown content found for %s", file_type)
                continue

            output_filename = f"{file_type.replace(' ', '_')}.docx"
            generated_file_path = await convert_markdown_to_docx(markdown_text, output_filename)
            if generated_file_path:
                generated_files.append(generated_file_path)

        logger.info("L1a generation completed in %.2fs", time.perf_counter() - total_start)
        return generated_files
    except Exception:
        logger.exception("Error in L1a document generation pipeline.")
        return []


def handle_doc_generation(file_manifests: list[dict], selected_options: list[str]) -> list[str]:
    """
    Synchronous wrapper that runs the async document generation pipeline.
    """
    return asyncio.run(async_handle_doc_generation(file_manifests, selected_options))
