import asyncio
import base64
import logging
import mimetypes
import os
import re
from typing import Any

import pypdfium2 as pdfium
from docx import Document as WordDocument
from pypdf import PdfReader

from .openai_client import get_openai_client


logger = logging.getLogger(__name__)

VISIBLE_CHARACTER_THRESHOLD = 80
WORD_THRESHOLD = 10

PDF_EXTENSIONS = {".pdf"}
WORD_EXTENSIONS = {".docx"}
TEXT_EXTENSIONS = {".txt", ".text", ".md", ".tex", ".csv", ".json", ".xml"}
IMAGE_EXTENSIONS = {".png", ".jpg", ".jpeg", ".bmp", ".tif", ".tiff", ".webp"}

BULLET_PATTERN = re.compile(r"^\s*(?:[-*•▪◦]+|\d+[.)]|[A-Za-z][.)])\s+")
WORD_PATTERN = re.compile(r"\b\w+\b")
FENCED_BLOCK_PATTERN = re.compile(r"```(?:latex)?\s*(.*?)```", re.DOTALL | re.IGNORECASE)
BULLET_PATTERN = re.compile(r"^\s*(?:[-*\u2022\u25AA\u25E6]+|\d+[.)]|[A-Za-z][.)])\s+")


def normalize_visible_text(text: str) -> str:
    return re.sub(r"\s+", " ", (text or "")).strip()


def count_visible_characters(text: str) -> int:
    return len(re.sub(r"\s+", "", text or ""))


def count_word_tokens(text: str) -> int:
    return len(WORD_PATTERN.findall(text or ""))


def is_scanned_page(text: str) -> bool:
    normalized = normalize_visible_text(text)
    return (
        count_visible_characters(normalized) < VISIBLE_CHARACTER_THRESHOLD
        or count_word_tokens(normalized) < WORD_THRESHOLD
    )


def latex_escape(value: str) -> str:
    replacements = {
        "\\": r"\textbackslash{}",
        "&": r"\&",
        "%": r"\%",
        "$": r"\$",
        "#": r"\#",
        "_": r"\_",
        "{": r"\{",
        "}": r"\}",
        "~": r"\textasciitilde{}",
        "^": r"\textasciicircum{}",
    }

    escaped = []
    for character in value or "":
        escaped.append(replacements.get(character, character))
    return "".join(escaped)


def clean_llm_latex_output(text: str) -> str:
    stripped = (text or "").strip()
    matches = FENCED_BLOCK_PATTERN.findall(stripped)
    if matches:
        stripped = "\n\n".join(match.strip() for match in matches if match.strip())
    return stripped.strip()


def line_looks_like_heading(line: str) -> bool:
    stripped = line.strip()
    if not stripped:
        return False
    if BULLET_PATTERN.match(stripped):
        return False
    return (
        stripped.isupper()
        or stripped.endswith(":")
        or (
            len(stripped) <= 80
            and sum(1 for token in stripped.split() if token[:1].isupper()) >= max(1, len(stripped.split()) - 1)
        )
    )


def lines_to_latex(lines: list[str]) -> str:
    blocks: list[str] = []
    index = 0

    while index < len(lines):
        stripped = lines[index].strip()
        if not stripped:
            index += 1
            continue

        if BULLET_PATTERN.match(stripped):
            items: list[str] = []
            ordered = bool(re.match(r"^\s*\d+[.)]\s+", stripped))

            while index < len(lines) and BULLET_PATTERN.match(lines[index].strip()):
                item = BULLET_PATTERN.sub("", lines[index].strip())
                items.append(latex_escape(item))
                index += 1

            env_name = "enumerate" if ordered else "itemize"
            blocks.append(f"\\begin{{{env_name}}}")
            blocks.extend(f"\\item {item}" for item in items)
            blocks.append(f"\\end{{{env_name}}}")
            continue

        if line_looks_like_heading(stripped):
            blocks.append(f"\\subsection*{{{latex_escape(stripped.rstrip(':'))}}}")
            index += 1
            continue

        paragraph_lines = [stripped]
        index += 1
        while index < len(lines):
            next_line = lines[index].strip()
            if not next_line or BULLET_PATTERN.match(next_line) or line_looks_like_heading(next_line):
                break
            paragraph_lines.append(next_line)
            index += 1

        blocks.append(latex_escape(" ".join(paragraph_lines)))

    return "\n\n".join(block for block in blocks if block).strip()


def text_to_latex(text: str) -> str:
    normalized = text.replace("\r\n", "\n").replace("\r", "\n")
    lines = [line.rstrip() for line in normalized.split("\n")]
    return lines_to_latex(lines)


def build_file_header(name: str, category: str, page_number: int) -> str:
    return "\n".join(
        [
            f"\\section*{{Source: {latex_escape(name)}}}",
            f"\\textbf{{Category}}: {latex_escape(category)}",
            f"\\textbf{{Page}}: {page_number}",
            "",
        ]
    ).strip()


def render_pdf_page_to_image(pdf_path: str, page_index: int, image_path: str) -> str:
    document = pdfium.PdfDocument(pdf_path)
    page = document[page_index]
    bitmap = page.render(scale=2)
    pil_image = bitmap.to_pil()
    pil_image.save(image_path)
    return image_path


def extract_docx_pages(file_path: str) -> list[dict[str, Any]]:
    document = WordDocument(file_path)
    blocks: list[str] = []
    in_list = False
    list_env = "itemize"

    for paragraph in document.paragraphs:
        text = paragraph.text.strip()
        if not text:
            if in_list:
                blocks.append(f"\\end{{{list_env}}}")
                in_list = False
            continue

        style_name = (paragraph.style.name or "").lower()
        if "heading" in style_name:
            if in_list:
                blocks.append(f"\\end{{{list_env}}}")
                in_list = False
            blocks.append(f"\\subsection*{{{latex_escape(text)}}}")
            continue

        if "list" in style_name or BULLET_PATTERN.match(text):
            ordered = "number" in style_name or bool(re.match(r"^\d+[.)]\s+", text))
            desired_env = "enumerate" if ordered else "itemize"
            if not in_list or desired_env != list_env:
                if in_list:
                    blocks.append(f"\\end{{{list_env}}}")
                blocks.append(f"\\begin{{{desired_env}}}")
                in_list = True
                list_env = desired_env

            item = BULLET_PATTERN.sub("", text)
            blocks.append(f"\\item {latex_escape(item)}")
            continue

        if in_list:
            blocks.append(f"\\end{{{list_env}}}")
            in_list = False

        blocks.append(latex_escape(text))

    if in_list:
        blocks.append(f"\\end{{{list_env}}}")

    return [
        {
            "page_number": 1,
            "latex_text": "\n\n".join(blocks).strip(),
            "extraction_mode": "docx_text",
        }
    ]


def extract_text_pages(file_path: str) -> list[dict[str, Any]]:
    with open(file_path, "r", encoding="utf-8", errors="ignore") as source_file:
        text = source_file.read()

    return [
        {
            "page_number": 1,
            "latex_text": text_to_latex(text),
            "extraction_mode": "plain_text",
        }
    ]


async def ocr_image_to_latex(image_path: str, source_name: str, page_number: int, ocr_model: str) -> str:
    mime_type = mimetypes.guess_type(image_path)[0] or "image/png"
    with open(image_path, "rb") as image_file:
        encoded_image = base64.b64encode(image_file.read()).decode("utf-8")

    prompt = (
        "You are a high-precision OCR assistant for immigration case files. "
        "Read the provided image carefully and convert the visible content into faithful LaTeX. "
        "Preserve headings, numbered lists, bullet lists, tables, labels, dates, and names exactly as shown. "
        "Do not summarize, correct, infer, or invent missing text. "
        "If a region is unreadable, write \\textit{[illegible]}. "
        "Return only LaTeX enclosed in triple backticks."
    )

    client = get_openai_client()
    response = await asyncio.to_thread(
        client.responses.create,
        model=ocr_model,
        input=[
            {
                "role": "user",
                "content": [
                    {"type": "input_text", "text": prompt},
                    {"type": "input_text", "text": f"Source: {source_name}. Page: {page_number}."},
                    {"type": "input_image", "image_url": f"data:{mime_type};base64,{encoded_image}"},
                ],
            }
        ],
    )
    return clean_llm_latex_output(response.output_text)


async def process_pdf_document(
    file_manifest: dict[str, Any],
    latex_output_dir: str,
    ocr_model: str,
) -> dict[str, Any]:
    reader = PdfReader(file_manifest["local_path"])
    page_entries: list[dict[str, Any]] = []
    extraction_modes: set[str] = set()

    for page_index, page in enumerate(reader.pages):
        extracted_text = page.extract_text() or ""
        page_number = page_index + 1

        if is_scanned_page(extracted_text):
            image_filename = f"{file_manifest['file_hash']}_page_{page_number}.png"
            image_path = os.path.join(latex_output_dir, image_filename)
            render_pdf_page_to_image(file_manifest["local_path"], page_index, image_path)
            latex_text = await ocr_image_to_latex(
                image_path=image_path,
                source_name=file_manifest["original_filename"],
                page_number=page_number,
                ocr_model=ocr_model,
            )
            extraction_mode = "vision_ocr"
        else:
            latex_text = text_to_latex(extracted_text)
            extraction_mode = "native_text"

        if not latex_text:
            latex_text = r"\textit{[no readable text extracted]}"

        extraction_modes.add(extraction_mode)
        page_entries.append(
            {
                "page_number": page_number,
                "latex_text": latex_text,
                "extraction_mode": extraction_mode,
            }
        )

    overall_mode = "hybrid" if len(extraction_modes) > 1 else next(iter(extraction_modes), "native_text")
    latex_path = os.path.join(latex_output_dir, f"{file_manifest['file_hash']}.tex")
    combined_latex = build_combined_latex(file_manifest, page_entries)
    with open(latex_path, "w", encoding="utf-8") as latex_file:
        latex_file.write(combined_latex)

    logger.info(
        "Processed PDF %s with %s pages using %s",
        file_manifest["original_filename"],
        len(page_entries),
        overall_mode,
    )

    return build_processed_manifest(file_manifest, latex_path, page_entries, overall_mode)


async def process_image_document(
    file_manifest: dict[str, Any],
    latex_output_dir: str,
    ocr_model: str,
) -> dict[str, Any]:
    latex_text = await ocr_image_to_latex(
        image_path=file_manifest["local_path"],
        source_name=file_manifest["original_filename"],
        page_number=1,
        ocr_model=ocr_model,
    )
    if not latex_text:
        latex_text = r"\textit{[no readable text extracted]}"

    page_entries = [
        {
            "page_number": 1,
            "latex_text": latex_text,
            "extraction_mode": "image_ocr",
        }
    ]
    latex_path = os.path.join(latex_output_dir, f"{file_manifest['file_hash']}.tex")
    with open(latex_path, "w", encoding="utf-8") as latex_file:
        latex_file.write(build_combined_latex(file_manifest, page_entries))

    logger.info("Processed image %s via OCR", file_manifest["original_filename"])
    return build_processed_manifest(file_manifest, latex_path, page_entries, "image_ocr")


async def process_non_pdf_document(file_manifest: dict[str, Any], latex_output_dir: str) -> dict[str, Any]:
    extension = (file_manifest.get("extension") or "").lower()
    if extension in WORD_EXTENSIONS:
        page_entries = extract_docx_pages(file_manifest["local_path"])
        overall_mode = "docx_text"
    elif extension in TEXT_EXTENSIONS:
        page_entries = extract_text_pages(file_manifest["local_path"])
        overall_mode = "plain_text"
    else:
        raise ValueError(f"Unsupported file extension for ingestion: {extension}")

    latex_path = os.path.join(latex_output_dir, f"{file_manifest['file_hash']}.tex")
    with open(latex_path, "w", encoding="utf-8") as latex_file:
        latex_file.write(build_combined_latex(file_manifest, page_entries))

    logger.info("Processed %s using %s", file_manifest["original_filename"], overall_mode)
    return build_processed_manifest(file_manifest, latex_path, page_entries, overall_mode)


def build_combined_latex(file_manifest: dict[str, Any], page_entries: list[dict[str, Any]]) -> str:
    sections: list[str] = []
    for page_entry in page_entries:
        sections.append(build_file_header(file_manifest["original_filename"], file_manifest["name"], page_entry["page_number"]))
        sections.append(page_entry["latex_text"])
        sections.append(r"\newpage")

    if sections and sections[-1] == r"\newpage":
        sections.pop()

    return "\n\n".join(section for section in sections if section).strip()


def build_processed_manifest(
    file_manifest: dict[str, Any],
    latex_path: str,
    page_entries: list[dict[str, Any]],
    overall_mode: str,
) -> dict[str, Any]:
    return {
        **file_manifest,
        "latex_path": latex_path,
        "page_count": len(page_entries),
        "extraction_mode": overall_mode,
        "page_entries": [
            {
                "page_number": page_entry["page_number"],
                "extraction_mode": page_entry["extraction_mode"],
                "latex_text": page_entry["latex_text"],
            }
            for page_entry in page_entries
        ],
    }


async def process_files_to_latex(
    file_manifests: list[dict[str, Any]],
    latex_output_dir: str,
    ocr_model: str,
) -> list[dict[str, Any]]:
    os.makedirs(latex_output_dir, exist_ok=True)
    processed_files: list[dict[str, Any]] = []

    for file_manifest in file_manifests:
        extension = (file_manifest.get("extension") or "").lower()
        if extension in PDF_EXTENSIONS:
            processed = await process_pdf_document(file_manifest, latex_output_dir, ocr_model)
        elif extension in IMAGE_EXTENSIONS:
            processed = await process_image_document(file_manifest, latex_output_dir, ocr_model)
        else:
            processed = await process_non_pdf_document(file_manifest, latex_output_dir)
        processed_files.append(processed)

    return processed_files
