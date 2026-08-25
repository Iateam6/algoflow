from __future__ import annotations

import os
from typing import Any

import markdown
from bs4 import BeautifulSoup
from bs4.element import NavigableString, Tag
from docx import Document


def _paragraph_ends_with_whitespace(paragraph) -> bool:
    runs = getattr(paragraph, "runs", None) or ()
    if not runs:
        return True
    text = runs[-1].text or ""
    return not text or text[-1].isspace()


def _append_text_with_breaks(paragraph, text: str, *, bold: bool = False) -> None:
    if text is None:
        return
    normalized = str(text).replace("\r\n", "\n").replace("\r", "\n")
    parts = normalized.split("\n")
    for index, part in enumerate(parts):
        if part:
            run = paragraph.add_run(part)
            if bold:
                run.bold = True
        if index < len(parts) - 1:
            paragraph.add_run().add_break()


def _append_inline_node(paragraph, node) -> None:
    if node is None:
        return
    if isinstance(node, Tag):
        if node.name == "br":
            paragraph.add_run().add_break()
            return
        if node.name == "strong":
            _append_text_with_breaks(paragraph, node.get_text(), bold=True)
            return
        _append_text_with_breaks(paragraph, node.get_text())
        return
    if isinstance(node, NavigableString):
        raw = str(node)
        if not raw:
            return
        if raw.isspace():
            if not _paragraph_ends_with_whitespace(paragraph):
                paragraph.add_run(" ")
            return
        _append_text_with_breaks(paragraph, raw)
        return


def markdown_to_docx(markdown_text: str, output_path: str) -> str:
    os.makedirs(os.path.dirname(output_path), exist_ok=True)
    soup = BeautifulSoup(
        markdown.markdown(markdown_text, extensions=["extra", "nl2br"]), "html.parser"
    )
    document = Document()
    for element in soup.children:
        if not isinstance(element, Tag):
            continue
        if element.name in {"h1", "h2", "h3"}:
            document.add_heading(element.get_text(), level=int(element.name[1]))
        elif element.name == "p":
            paragraph = document.add_paragraph()
            for node in element.children:
                _append_inline_node(paragraph, node)
        elif element.name == "hr":
            document.add_page_break()
        elif element.name == "table":
            rows = element.find_all("tr")
            column_count = max(
                (len(row.find_all(["th", "td"], recursive=False)) for row in rows),
                default=0,
            )
            if column_count:
                table = document.add_table(rows=len(rows), cols=column_count)
                table.style = "Table Grid"
                for row_index, source_row in enumerate(rows):
                    source_cells = source_row.find_all(["th", "td"], recursive=False)
                    for column_index, source_cell in enumerate(source_cells):
                        paragraph = table.rows[row_index].cells[column_index].paragraphs[0]
                        if source_cell.name == "th":
                            paragraph.add_run(source_cell.get_text(" ", strip=True)).bold = True
                        else:
                            for node in source_cell.children:
                                _append_inline_node(paragraph, node)
        elif element.name in {"ul", "ol"}:
            style = "ListBullet" if element.name == "ul" else "ListNumber"
            for item in element.find_all("li", recursive=False):
                document.add_paragraph(item.get_text(), style=style)
    document.save(output_path)
    return output_path

