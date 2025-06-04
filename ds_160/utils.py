import os
import re
import openai
import asyncio
import logging
import aiohttp
import aiofiles
import textwrap
import markdown
from io import BytesIO
from openai import OpenAI
from docx import Document
from bs4 import BeautifulSoup
from .agent import initialize_agents ,generate_document
from dotenv import load_dotenv
from django.conf import settings

# Load environment variables
load_dotenv()
print("API Key:", os.getenv("OPENAI_API_KEY"))

# Configure logging
logging.basicConfig(level=logging.INFO)

# Initialize OpenAI client
openai.api_key = os.getenv("OPENAI_API_KEY")

client = OpenAI(
    api_key=os.getenv("OPENAI_API_KEY")
)

VECTOR_STORE_ID = None

async def create_file(client, file_path):
    if file_path.startswith("http://") or file_path.startswith("https://"):
        # Download the file content from the URL asynchronously
        async with aiohttp.ClientSession() as session:
            async with session.get(file_path) as response:
                file_content = BytesIO(await response.read())
                file_name = file_path.split("/")[-1]
                file_tuple = (file_name, file_content)
                result = await client.files.create(
                    file=file_tuple,
                    purpose="assistants"
                )
    else:
        # Handle local file path asynchronously
        async with aiofiles.open(file_path, "rb") as file:
            file_content = await file.read()  # Read the file content as bytes
            file_name = os.path.basename(file_path)
            file_tuple = (file_name, BytesIO(file_content))
            result = client.files.create(
                file=file_tuple,
                purpose="assistants"
            )
    print(result.id)
    return result.id

async def upload_files(file_paths):
    """
    Upload files to the vector store. Reuse the vector store if it already exists.
    """
    global VECTOR_STORE_ID

    # Reuse the existing vector store if available
    if VECTOR_STORE_ID:
        logging.info(f"Reusing existing vector store: {VECTOR_STORE_ID}")
        return VECTOR_STORE_ID

    # Create a new vector store
    vector_store = client.vector_stores.create(name="DS-160 (Nonimmigrant Visas) store")
    VECTOR_STORE_ID = vector_store.id
    logging.info(f"Created new vector store: {VECTOR_STORE_ID}")

    # Upload files concurrently
    tasks = [create_file(client, path) for path in file_paths]
    file_ids = await asyncio.gather(*tasks)

    # Add files to the vector store
    for file_id in file_ids:
        client.vector_stores.files.create(
            vector_store_id=VECTOR_STORE_ID,
            file_id=file_id
        )
        logging.info(f"Added file {file_id} to vector store {VECTOR_STORE_ID}")

    return VECTOR_STORE_ID

async def generate_documents(requested_files, agents):
    """
    Generate the requested documents sequentially using the corresponding agents.

    Args:
        requested_files (list): A list of document types to generate (e.g., ["petition_cover_letter", "employer_support_letter"]).
        agents (dict): Dictionary of initialized agents.
    """
    for file_type in requested_files:
        agent_response = await generate_document(file_type, agents)
        if agent_response:
            # Extract Markdown content
            matches = re.findall(r"```(.*?)```", str(agent_response), re.DOTALL)
            if not matches:
                logging.warning(f"No Markdown content found for {file_type}")
                continue

            for i, match in enumerate(matches, 1):
                documents_text = textwrap.dedent(match).strip()
                logging.info(f"Match {i} for {file_type}: {documents_text[:100]}...")  # Log first 100 chars

                # Convert the generated Markdown to .docx using the agent's name as the file name
                output_file = f"{file_type.replace(' ', '_')}.docx"
                try:
                    await convert_markdown_to_docx(documents_text, output_file)
                    logging.info(f"Generated {output_file}")
                except Exception as e:
                    logging.error(f"Error generating {output_file}: {e}")
        else:
            logging.warning(f"Failed to generate document for {file_type}")

# Convert Markdown to .docx
async def convert_markdown_to_docx(md_text: str, output_path: str) -> str:
    """
    Convert Markdown text to a .docx file asynchronously and save it in the 'media/generated' folder.
    If files already exist in the directory, delete them before saving the new file.

    Args:
        md_text (str): The Markdown text to convert.
        output_path (str): The name of the file to save (not the full path).

    Returns:
        str: The absolute path to the generated .docx file.
    """
    try:
        # Use the 'generated' directory inside MEDIA_ROOT
        generated_dir = os.path.join(settings.MEDIA_ROOT, "generated")
        os.makedirs(generated_dir, exist_ok=True)  # Ensure the directory exists

        # Check and delete existing files in the directory
        for filename in os.listdir(generated_dir):
            file_path = os.path.join(generated_dir, filename)
            try:
                if os.path.isfile(file_path):
                    os.remove(file_path)
                    logging.info(f"Deleted existing file: {file_path}")
            except OSError as e:
                logging.error(f"Error deleting file {file_path}: {e}")

        # Full path for the output file
        full_output_path = os.path.join(generated_dir, output_path)

        # Convert Markdown to HTML
        html = markdown.markdown(md_text, extensions=['extra'])
        soup = BeautifulSoup(html, 'html.parser')

        # Create a new Word document
        doc = Document()

        # Iterate over top-level HTML elements
        for element in soup.children:
            if element.name == 'h1':
                doc.add_heading(element.get_text(), level=1)
            elif element.name == 'h2':
                doc.add_heading(element.get_text(), level=2)
            elif element.name == 'h3':
                doc.add_heading(element.get_text(), level=3)
            elif element.name == 'p':
                para = doc.add_paragraph()
                for node in element.children:
                    text = node.get_text() if hasattr(node, 'get_text') else str(node)
                    if node.name == 'strong':
                        run = para.add_run(text)
                        run.bold = True
                    else:
                        para.add_run(text)
            elif element.name == 'hr':
                doc.add_page_break()
            elif element.name == 'ul':
                for li in element.find_all('li', recursive=False):
                    doc.add_paragraph(li.get_text(), style='ListBullet')
            elif element.name == 'ol':
                for li in element.find_all('li', recursive=False):
                    doc.add_paragraph(li.get_text(), style='ListNumber')
            elif isinstance(element, str) and element.strip():
                doc.add_paragraph(element.strip())

        # Save the .docx file
        doc.save(full_output_path)
        logging.info(f"Saved .docx file to {full_output_path}")
        return full_output_path
    except Exception as e:
        logging.error(f"Error converting Markdown to .docx: {e}")
        return None

async def async_handle_doc_generation(file_paths, selected_options):
    """
    Generate documents based on the uploaded files and selected options.

    Args:
        file_paths (list): List of file paths to process.
        selected_options (list): List of selected options for document generation.

    Returns:
        list: Paths to the generated documents.
    """
    try:
        # Upload files to the vector store
        vector_store_id = await upload_files(file_paths)

        # Initialize agents with the vector store ID
        agents = await initialize_agents(vector_store_id)

        # Generate the requested documents
        generated_files = []
        for file_type in selected_options:
            agent_response = await generate_document(file_type, agents)
            if agent_response:
                # Extract Markdown content
                extracted_text = re.findall(r"```(.*?)```", str(agent_response), re.DOTALL)
                if extracted_text:
                    # Join the extracted text into a single string (if multiple blocks exist)
                    documents_text = "\n".join(textwrap.dedent(match).strip() for match in extracted_text)
                    logging.info(f"Extracted text for {file_type}: {documents_text[:100]}...")  # Log first 100 chars

                    # Convert Markdown to .docx and append the file path
                    output_file_name = f"{file_type.replace(' ', '_')}.docx"
                    generated_file_path = await convert_markdown_to_docx(documents_text, output_file_name)
                    if generated_file_path:
                        generated_files.append(generated_file_path)
                    logging.info(f"Generated {generated_file_path}")
                else:
                    logging.warning(f"No Markdown content found for {file_type}")
            else:
                logging.warning(f"No response from agent for {file_type}")

        return generated_files
    except Exception as e:
        logging.error(f"Error in document generation: {e}")
        return []

# Synchronous wrapper for the async function
def handle_doc_generation(file_paths, selected_options):
    """
    Sync wrapper that runs the async document generation pipeline.
    """
    return asyncio.run(async_handle_doc_generation(file_paths, selected_options))

