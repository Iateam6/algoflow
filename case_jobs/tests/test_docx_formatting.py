import asyncio
import tempfile

from django.test import SimpleTestCase, override_settings
from docx import Document

from case_jobs.pipeline.legacy_generation import convert_markdown_to_docx


class DocxFormattingTests(SimpleTestCase):
    def test_converter_renders_markdown_tables(self):
        md_text = """| Exhibit | Supporting Document |
|---|---|
| **Exhibit 1** | **Forms and Fees** |
| | 1.1 Form G-28 |
"""
        with tempfile.TemporaryDirectory(dir="media") as tmpdir:
            with override_settings(MEDIA_ROOT=tmpdir):
                output_path = asyncio.run(convert_markdown_to_docx(md_text, "table.docx"))

            document = Document(output_path)
            self.assertEqual(len(document.tables), 1)
            self.assertEqual(document.tables[0].cell(0, 0).text, "Exhibit")
            self.assertIn("1.1 Form G-28", document.tables[0].cell(2, 1).text)

    def test_converter_preserves_single_newlines_and_run_spacing(self):
        md_text = "\n\n".join(
            (
                "Employer: X\nBeneficiary: Y\nPosition: Z\nCountry: W",
                "Sincerely,\n**Alex Smith**\nLakshmi",
            )
        )
        with tempfile.TemporaryDirectory(dir="media") as tmpdir:
            with override_settings(MEDIA_ROOT=tmpdir):
                output_path = asyncio.run(convert_markdown_to_docx(md_text, "out.docx"))

            document = Document(output_path)
            paragraphs = [paragraph.text for paragraph in document.paragraphs if paragraph.text]
            self.assertGreaterEqual(len(paragraphs), 2)

            header = paragraphs[0]
            self.assertIn("Employer: X", header)
            self.assertIn("\nBeneficiary: Y", header)
            self.assertIn("\nPosition: Z", header)
            self.assertIn("\nCountry: W", header)

            signature = "\n".join(paragraphs[1:])
            self.assertNotIn("Alex SmithLakshmi", signature)
            self.assertIn("Alex Smith", signature)
            self.assertIn("Lakshmi", signature)
            self.assertIn("\n", signature)
